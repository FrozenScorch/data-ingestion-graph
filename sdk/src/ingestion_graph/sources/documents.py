"""Checkpoint-safe local document and folder source."""

from __future__ import annotations

import asyncio
import csv
import hashlib
import json
import mimetypes
import os
import stat
import tempfile
import zipfile
from collections.abc import AsyncIterator, Iterable, Iterator, Mapping, Sequence
from contextlib import contextmanager, suppress
from dataclasses import dataclass
from datetime import UTC, date, datetime, time, timedelta
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from ingestion_graph.connectors.base import (
    CheckResult,
    ConnectorCapabilities,
    ConnectorSpec,
    Source,
    StreamDescriptor,
)
from ingestion_graph.document_ai import canonical_fingerprint, evaluate_text_quality
from ingestion_graph.errors import ConfigurationError
from ingestion_graph.messages import LogMessage, RecordMessage, SourceMessage, StateMessage
from ingestion_graph.models import (
    DocumentElement,
    Envelope,
    Operation,
    Payload,
    TableBatch,
    Tombstone,
    stable_record_id,
)

WordDocument: Any
try:
    from docx import Document as _WordDocument

    WordDocument = _WordDocument
except ImportError:  # pragma: no cover - exercised without the optional extra
    WordDocument = None

load_workbook: Any
try:
    from openpyxl import load_workbook as _load_workbook  # type: ignore[import-untyped]

    load_workbook = _load_workbook
except ImportError:  # pragma: no cover - exercised without the optional extra
    load_workbook = None

PdfReader: Any
try:
    from pypdf import PdfReader as _PdfReader

    PdfReader = _PdfReader
except ImportError:  # pragma: no cover - exercised without the optional extra
    PdfReader = None


SUPPORTED_EXTENSIONS = (
    ".csv",
    ".docx",
    ".eml",
    ".htm",
    ".html",
    ".md",
    ".pdf",
    ".txt",
    ".xlsx",
)
OCR_IMAGE_EXTENSIONS = (".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp")
ALL_DOCUMENT_EXTENSIONS = SUPPORTED_EXTENSIONS + OCR_IMAGE_EXTENSIONS
PARSER_VERSION = "2"


@dataclass(frozen=True, slots=True)
class _ParsedElement:
    payload: Payload
    metadata: Mapping[str, Any]


@dataclass(frozen=True, slots=True)
class _Snapshot:
    path: Path
    sha256: str
    stat: os.stat_result


@dataclass(frozen=True, slots=True)
class _TrustedRoot:
    path: Path
    identity: tuple[int, int]
    is_file: bool


class LocalDocumentsSource(Source):
    """Read supported files or directory trees as canonical SDK envelopes.

    Each configured root is a stream. Per-file SHA-256 and element-count state
    supports mid-document resume, changed-file reconciliation, and tombstones
    when files disappear from a directory root.
    """

    def __init__(
        self,
        paths: str | Path | Sequence[str | Path],
        *,
        recursive: bool = True,
        extensions: Sequence[str] = SUPPORTED_EXTENSIONS,
        checkpoint_interval: int = 50,
        text_chunk_chars: int = 12_000,
        table_batch_rows: int = 500,
        include_hidden: bool = False,
        follow_symlinks: bool = False,
        stream_names: Sequence[str] | None = None,
        max_file_size_bytes: int = 256 * 1024 * 1024,
        max_archive_uncompressed_bytes: int = 512 * 1024 * 1024,
        ocr_mode: str = "off",
        table_mode: str = "off",
        ocr_engine: Any = None,
        page_renderer: Any = None,
        table_extractor: Any = None,
        document_splitter: Any = None,
        extraction_cache: Any = None,
        failure_mode: str = "strict",
        min_native_text_quality: float = 0.65,
        vision_extractor: Any = None,
        vision_fallback: bool = False,
        external_processing_policy: Any = None,
        ocr_languages: Sequence[str] = ("eng",),
        render_dpi: int = 300,
        page_timeout_seconds: float = 120.0,
        max_page_concurrency: int = 2,
        retain_extraction_artifacts: bool = False,
    ) -> None:
        raw_paths = (paths,) if isinstance(paths, (str, Path)) else tuple(paths)
        if not raw_paths:
            raise ConfigurationError("Document paths must not be empty")
        normalized_extensions = tuple(_normalize_extension(item) for item in extensions)
        if not normalized_extensions:
            raise ConfigurationError("Document extensions must not be empty")
        supported = set(ALL_DOCUMENT_EXTENSIONS)
        unsupported = sorted(set(normalized_extensions) - supported)
        if unsupported:
            raise ConfigurationError(f"Unsupported document extensions: {', '.join(unsupported)}")
        if (
            any(item in OCR_IMAGE_EXTENSIONS for item in normalized_extensions)
            and ocr_mode == "off"
        ):
            raise ConfigurationError("Image extensions require ocr_mode='auto' or 'always'")
        if ocr_mode not in {"off", "auto", "always"}:
            raise ConfigurationError("Document ocr_mode must be off, auto, or always")
        if table_mode not in {"off", "native", "auto", "local", "vision"}:
            raise ConfigurationError("Document table_mode is invalid")
        if table_mode == "vision" and vision_extractor is None:
            raise ConfigurationError("Document table_mode='vision' requires vision_extractor")
        if vision_fallback and vision_extractor is None:
            raise ConfigurationError("vision_fallback requires vision_extractor")
        if vision_extractor is not None and external_processing_policy is None:
            raise ConfigurationError("vision_extractor requires external_processing_policy")
        if failure_mode not in {"strict", "best_effort"}:
            raise ConfigurationError("Document failure_mode must be strict or best_effort")
        if not isinstance(min_native_text_quality, (int, float)) or isinstance(
            min_native_text_quality, bool
        ):
            raise ConfigurationError("Document min_native_text_quality must be numeric")
        if not 0 <= float(min_native_text_quality) <= 1:
            raise ConfigurationError("Document min_native_text_quality must be between 0 and 1")
        if not ocr_languages or any(
            not isinstance(item, str) or not item or not item.replace("-", "").isalnum()
            for item in ocr_languages
        ):
            raise ConfigurationError("Document ocr_languages must contain valid language codes")
        if isinstance(render_dpi, bool) or render_dpi <= 0:
            raise ConfigurationError("Document render_dpi must be positive")
        if (
            isinstance(page_timeout_seconds, bool)
            or not isinstance(page_timeout_seconds, (int, float))
            or page_timeout_seconds <= 0
        ):
            raise ConfigurationError("Document page_timeout_seconds must be positive")
        if isinstance(max_page_concurrency, bool) or max_page_concurrency <= 0:
            raise ConfigurationError("Document max_page_concurrency must be positive")
        if checkpoint_interval < 1:
            raise ConfigurationError("Document checkpoint_interval must be positive")
        if text_chunk_chars < 256:
            raise ConfigurationError("Document text_chunk_chars must be at least 256")
        if table_batch_rows < 1:
            raise ConfigurationError("Document table_batch_rows must be positive")
        if max_file_size_bytes < 1 or max_archive_uncompressed_bytes < 1:
            raise ConfigurationError("Document size limits must be positive")
        if stream_names is not None:
            normalized_names = tuple(str(item).strip() for item in stream_names)
            if len(normalized_names) != len(raw_paths) or any(
                not item for item in normalized_names
            ):
                raise ConfigurationError("Document stream_names must name every configured path")
            if len(set(normalized_names)) != len(normalized_names):
                raise ConfigurationError("Document stream_names must be unique")
        else:
            normalized_names = ()

        self.paths = tuple(Path(item).expanduser() for item in raw_paths)
        self.recursive = recursive
        self.extensions = normalized_extensions
        self.checkpoint_interval = checkpoint_interval
        self.text_chunk_chars = text_chunk_chars
        self.table_batch_rows = table_batch_rows
        self.include_hidden = include_hidden
        self.follow_symlinks = follow_symlinks
        self.stream_names = normalized_names
        self.max_file_size_bytes = max_file_size_bytes
        self.max_archive_uncompressed_bytes = max_archive_uncompressed_bytes
        self.ocr_mode = ocr_mode
        self.table_mode = table_mode
        self.ocr_engine = ocr_engine
        self.page_renderer = page_renderer
        self.table_extractor = table_extractor
        self.document_splitter = document_splitter
        self.extraction_cache = extraction_cache
        self.failure_mode = failure_mode
        self.min_native_text_quality = float(min_native_text_quality)
        self.vision_extractor = vision_extractor
        self.vision_fallback = vision_fallback
        self.external_processing_policy = external_processing_policy
        self.ocr_languages = tuple(ocr_languages)
        self.render_dpi = int(render_dpi)
        self.page_timeout_seconds = float(page_timeout_seconds)
        self.max_page_concurrency = int(max_page_concurrency)
        self.retain_extraction_artifacts = bool(retain_extraction_artifacts)
        active_components = [self.document_splitter]
        if self.ocr_mode != "off":
            active_components.extend((self.ocr_engine, self.page_renderer))
        if self.table_mode in {"auto", "local"}:
            active_components.extend((self.table_extractor, self.page_renderer))
        if self.table_mode == "vision" or self.vision_fallback:
            active_components.extend((self.vision_extractor, self.page_renderer))
        for component in active_components:
            descriptor = getattr(component, "descriptor", None)
            if (
                descriptor is not None
                and (descriptor.external or not descriptor.deterministic)
                and (extraction_cache is None or not getattr(extraction_cache, "persistent", False))
            ):
                raise ConfigurationError(
                    "External or nondeterministic document components require a persistent "
                    "extraction cache"
                )
        self._runtime_ocr_engine: Any = None
        self._runtime_page_renderer: Any = None
        self._runtime_table_extractor: Any = None
        self._closed = False
        self._streams: dict[str, Path] = {}

    @classmethod
    def manifest(cls) -> ConnectorSpec:
        return ConnectorSpec(
            name="local_documents",
            version="1.0.0",
            config_schema={
                "type": "object",
                "properties": {
                    "paths": {
                        "type": "array",
                        "items": {"type": "string"},
                        "minItems": 1,
                        "format": "local-paths",
                    },
                    "recursive": {"type": "boolean", "default": True},
                    "extensions": {
                        "type": "array",
                        "items": {"type": "string", "enum": list(ALL_DOCUMENT_EXTENSIONS)},
                        "default": list(SUPPORTED_EXTENSIONS),
                    },
                    "checkpoint_interval": {
                        "type": "integer",
                        "minimum": 1,
                        "default": 50,
                    },
                    "text_chunk_chars": {
                        "type": "integer",
                        "minimum": 256,
                        "default": 12_000,
                    },
                    "table_batch_rows": {
                        "type": "integer",
                        "minimum": 1,
                        "default": 500,
                    },
                    "include_hidden": {"type": "boolean", "default": False},
                    "follow_symlinks": {"type": "boolean", "default": False},
                    "stream_names": {
                        "type": "array",
                        "items": {"type": "string", "minLength": 1},
                        "description": "Optional stable stream name for each configured path",
                    },
                    "max_file_size_bytes": {
                        "type": "integer",
                        "minimum": 1,
                        "default": 268_435_456,
                    },
                    "max_archive_uncompressed_bytes": {
                        "type": "integer",
                        "minimum": 1,
                        "default": 536_870_912,
                    },
                    "ocr_mode": {
                        "type": "string",
                        "enum": ["off", "auto", "always"],
                        "default": "off",
                    },
                    "table_mode": {
                        "type": "string",
                        "enum": ["off", "native", "auto", "local", "vision"],
                        "default": "off",
                    },
                    "failure_mode": {
                        "type": "string",
                        "enum": ["strict", "best_effort"],
                        "default": "strict",
                    },
                    "min_native_text_quality": {
                        "type": "number",
                        "minimum": 0,
                        "maximum": 1,
                        "default": 0.65,
                    },
                    "ocr_languages": {"type": "array", "items": {"type": "string"}},
                    "render_dpi": {"type": "integer", "minimum": 1, "default": 300},
                    "page_timeout_seconds": {"type": "number", "exclusiveMinimum": 0},
                    "max_page_concurrency": {"type": "integer", "minimum": 1},
                    "vision_fallback": {"type": "boolean", "default": False},
                },
                "required": ["paths"],
            },
            capabilities=ConnectorCapabilities(
                incremental=True,
                resumable_full_refresh=True,
                deletes=True,
                schema_discovery=True,
            ),
        )

    def spec(self) -> ConnectorSpec:
        return self.manifest()

    async def check(self) -> CheckResult:
        missing = [str(path) for path in self.paths if not path.exists()]
        invalid_missing = [
            value for value in missing if Path(value).suffix.lower() not in self.extensions
        ]
        if invalid_missing:
            return CheckResult(False, f"Document paths do not exist: {', '.join(invalid_missing)}")
        if not self.follow_symlinks and any(path.is_symlink() for path in self.paths):
            return CheckResult(False, "Configured document roots must not be symlinks")
        try:
            files = self._discover_files()
        except OSError as exc:
            return CheckResult(False, str(exc))
        dependency_error = _missing_dependency(files)
        if dependency_error:
            return CheckResult(False, dependency_error)
        if self.ocr_mode != "off":
            engine = self.ocr_engine
            if engine is None:
                from ingestion_graph.document_ai.tesseract import TesseractOcrEngine

                engine = TesseractOcrEngine()
            engine_check = await engine.check()
            check = CheckResult(bool(engine_check.ok), str(engine_check.message))
            if not check.ok:
                return check
        oversized = [str(path) for path in files if path.stat().st_size > self.max_file_size_bytes]
        if oversized:
            return CheckResult(
                False, f"Document files exceed the configured size limit: {oversized[0]}"
            )
        message = f"Discovered {len(files)} supported document files"
        if missing:
            message += f"; missing roots will be reconciled as deletions: {', '.join(missing)}"
        return CheckResult(True, message)

    async def discover(self) -> Sequence[StreamDescriptor]:
        self._streams = self._name_streams()
        return [
            StreamDescriptor(
                name=name,
                namespace="local.documents",
                primary_key=("element_id",),
                cursor_field=("files", "in_progress"),
                json_schema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "filename": {"type": "string"},
                        "extension": {"type": "string"},
                        "relative_path": {"type": "string"},
                        "element_index": {"type": "integer"},
                    },
                },
            )
            for name in self._streams
        ]

    async def read(
        self,
        stream: StreamDescriptor,
        state: Mapping[str, Any] | None = None,
    ) -> AsyncIterator[SourceMessage]:
        if stream.name not in self._streams:
            await self.discover()
        root = self._streams.get(stream.name)
        if root is None:
            raise ConfigurationError(f"Document stream {stream.name!r} is not configured")

        current = _validate_state(state)
        files_state = dict(current["files"])
        parser_fingerprint = self._parser_fingerprint()
        in_progress = current.get("in_progress")
        pending_in_progress = dict(in_progress) if isinstance(in_progress, Mapping) else None
        trusted_root = _trusted_root(root) if not self.follow_symlinks else None
        discovered = self._discover_files(root)
        current_paths = {_relative_path(path, root): path for path in discovered}

        removed_paths = set(files_state) - set(current_paths)
        if isinstance(in_progress, Mapping):
            progress_path = in_progress.get("relative_path")
            if isinstance(progress_path, str) and progress_path not in current_paths:
                removed_paths.add(progress_path)
        progress_path = (
            pending_in_progress.get("relative_path") if pending_in_progress is not None else None
        )
        ordered_removed_paths = sorted(
            removed_paths,
            key=lambda item: (0 if item == progress_path else 1, item),
        )
        for relative_path in ordered_removed_paths:
            prior = files_state.get(relative_path)
            prior_count = int(prior["element_count"]) if prior is not None else 0
            progress_count = (
                max(
                    int(in_progress.get("next_index", 0)),
                    int(in_progress.get("reconcile_element_count", 0)),
                )
                if isinstance(in_progress, Mapping)
                and in_progress.get("relative_path") == relative_path
                else 0
            )
            delete_count = max(prior_count, progress_count)
            if pending_in_progress is None and prior is not None:
                pending_in_progress = {
                    "relative_path": relative_path,
                    "sha256": str(prior["sha256"]),
                    "next_index": delete_count,
                    "parser_fingerprint": prior.get("parser_fingerprint") or parser_fingerprint,
                    "tombstone_next_index": 0,
                }
                for key in ("manifest_cache_key", "manifest_sha256"):
                    if prior.get(key) is not None:
                        pending_in_progress[key] = prior[key]
            elif (
                pending_in_progress is not None
                and pending_in_progress.get("relative_path") == relative_path
            ):
                pending_in_progress["next_index"] = delete_count
                if pending_in_progress.get("reconcile_element_count") is not None:
                    pending_in_progress["reconcile_element_count"] = delete_count
            tombstone_start = (
                int(pending_in_progress.get("tombstone_next_index", 0))
                if pending_in_progress is not None
                and pending_in_progress.get("relative_path") == relative_path
                else 0
            )
            async for message in self._emit_tombstones(
                stream.name,
                relative_path,
                delete_count,
                files_state,
                parser_fingerprint,
                start=tombstone_start,
                in_progress=pending_in_progress,
            ):
                yield message
            files_state.pop(relative_path, None)
            if (
                pending_in_progress is not None
                and pending_in_progress.get("relative_path") == relative_path
            ):
                pending_in_progress = None
            yield StateMessage(
                stream.name,
                _checkpoint(files_state, parser_fingerprint, pending_in_progress),
            )

        progress_path = (
            pending_in_progress.get("relative_path") if pending_in_progress is not None else None
        )
        ordered_current_paths = sorted(
            current_paths.items(),
            key=lambda item: (0 if item[0] == progress_path else 1, item[0]),
        )
        for relative_path, path in ordered_current_paths:
            prior = files_state.get(relative_path)
            with _snapshot_file(
                path,
                root=root,
                trusted_root=trusted_root,
                follow_symlinks=self.follow_symlinks,
                max_bytes=self.max_file_size_bytes,
            ) as snapshot:
                fingerprint = snapshot.sha256
                if (
                    prior is not None
                    and prior.get("parser_fingerprint") == parser_fingerprint
                    and prior["sha256"] == fingerprint
                    and not (
                        isinstance(in_progress, Mapping)
                        and in_progress.get("relative_path") == relative_path
                    )
                ):
                    continue
                if path.suffix.lower() in {".docx", ".xlsx"}:
                    _validate_archive(
                        snapshot.path,
                        max_uncompressed_bytes=self.max_archive_uncompressed_bytes,
                    )
                resume_index = 0
                if (
                    isinstance(in_progress, Mapping)
                    and in_progress.get("relative_path") == relative_path
                    and in_progress.get("sha256") == fingerprint
                    and in_progress.get("parser_fingerprint") == parser_fingerprint
                    and (
                        in_progress.get("tombstone_next_index") is None
                        or in_progress.get("reconcile_element_count") is not None
                    )
                ):
                    resume_index = int(in_progress.get("next_index", 0))

                prior_count = int(prior["element_count"]) if prior is not None else 0
                state_progress_count = (
                    int(in_progress.get("next_index", 0))
                    if isinstance(in_progress, Mapping)
                    and in_progress.get("relative_path") == relative_path
                    else 0
                )
                reconcile_count = (
                    int(in_progress.get("reconcile_element_count", 0))
                    if isinstance(in_progress, Mapping)
                    and in_progress.get("relative_path") == relative_path
                    else 0
                )
                committed_count = max(prior_count, state_progress_count, reconcile_count)
                reconciliation_active = bool(
                    committed_count
                    and (
                        reconcile_count
                        or (
                            prior is not None
                            and (
                                prior.get("sha256") != fingerprint
                                or prior.get("parser_fingerprint") != parser_fingerprint
                            )
                        )
                        or (
                            isinstance(in_progress, Mapping)
                            and in_progress.get("relative_path") == relative_path
                            and (
                                in_progress.get("sha256") != fingerprint
                                or in_progress.get("parser_fingerprint") != parser_fingerprint
                            )
                        )
                    )
                )

                parsed_count = 0
                emitted = 0
                manifest_key: str | None = None
                manifest_sha256: str | None = None
                try:
                    elements: list[_ParsedElement] | None = None
                    requires_manifest = self._requires_replay_manifest()
                    checkpoint_matches_snapshot = bool(
                        isinstance(in_progress, Mapping)
                        and in_progress.get("relative_path") == relative_path
                        and in_progress.get("sha256") == fingerprint
                        and in_progress.get("parser_fingerprint") == parser_fingerprint
                    )
                    referenced_manifest = (
                        in_progress.get("manifest_cache_key")
                        if checkpoint_matches_snapshot and isinstance(in_progress, Mapping)
                        else None
                    )
                    requires_cached_manifest = bool(
                        requires_manifest
                        and checkpoint_matches_snapshot
                        and isinstance(in_progress, Mapping)
                        and (
                            resume_index > 0 or in_progress.get("tombstone_next_index") is not None
                        )
                    )
                    if requires_manifest:
                        if self.extraction_cache is None or not getattr(
                            self.extraction_cache, "persistent", False
                        ):
                            raise ConfigurationError(
                                "Nondeterministic document components require a persistent "
                                "extraction cache"
                            )
                        manifest_key = canonical_fingerprint(
                            {
                                "kind": "ordered-extraction-manifest",
                                "schema_version": "1",
                                "document_sha256": fingerprint,
                                "source_extension": snapshot.path.suffix.lower(),
                                "parser_fingerprint": parser_fingerprint,
                            }
                        )
                        if referenced_manifest is not None and referenced_manifest != manifest_key:
                            raise ConfigurationError(
                                "In-progress extraction manifest does not match this document"
                            )
                        cached_manifest = await self.extraction_cache.get(manifest_key)
                        if cached_manifest is not None:
                            try:
                                expected_hash = (
                                    in_progress.get("manifest_sha256")
                                    if checkpoint_matches_snapshot
                                    and isinstance(in_progress, Mapping)
                                    else None
                                )
                                manifest_sha256 = hashlib.sha256(cached_manifest).hexdigest()
                                if expected_hash is not None and expected_hash != manifest_sha256:
                                    raise ValueError("manifest hash mismatch")
                                elements = _deserialize_extraction_manifest(cached_manifest)
                            except (TypeError, ValueError, json.JSONDecodeError) as exc:
                                if referenced_manifest is not None or requires_cached_manifest:
                                    raise ConfigurationError(
                                        "In-progress extraction manifest is corrupt"
                                    ) from exc
                                await self.extraction_cache.delete(manifest_key)
                                elements = None
                        elif referenced_manifest is not None or requires_cached_manifest:
                            raise ConfigurationError(
                                "In-progress extraction manifest is unavailable"
                            )
                    if elements is None:
                        if (
                            self.ocr_mode == "off"
                            and self.table_mode == "off"
                            and self.document_splitter is None
                        ):
                            elements = list(
                                _parse_file(
                                    snapshot.path,
                                    text_chunk_chars=self.text_chunk_chars,
                                    table_batch_rows=self.table_batch_rows,
                                )
                            )
                        else:
                            elements = await self._parse_file_with_document_ai(snapshot.path)
                        if manifest_key is not None:
                            manifest_bytes = _serialize_extraction_manifest(elements)
                            await self.extraction_cache.put(manifest_key, manifest_bytes)
                            manifest_sha256 = hashlib.sha256(manifest_bytes).hexdigest()
                except asyncio.CancelledError:
                    raise
                except Exception as exc:
                    if self.failure_mode == "strict":
                        raise
                    yield LogMessage(
                        "warning",
                        "document_ai_file_failed",
                        {"extension": path.suffix.lower(), "error_type": type(exc).__name__},
                    )
                    continue
                for index, element in enumerate(elements):
                    parsed_count = index + 1
                    if index < resume_index:
                        continue
                    yield RecordMessage(
                        self._record_envelope(
                            stream.name,
                            path,
                            relative_path,
                            fingerprint,
                            snapshot.stat,
                            index,
                            element,
                        )
                    )
                    emitted += 1
                    if emitted >= self.checkpoint_interval:
                        if (
                            pending_in_progress is not None
                            and pending_in_progress.get("relative_path") != relative_path
                        ):
                            emitted = 0
                            continue
                        progress: dict[str, Any] = {
                            "relative_path": relative_path,
                            "sha256": fingerprint,
                            "next_index": index + 1,
                            "parser_fingerprint": parser_fingerprint,
                        }
                        if manifest_key is not None and manifest_sha256 is not None:
                            progress.update(
                                {
                                    "manifest_cache_key": manifest_key,
                                    "manifest_sha256": manifest_sha256,
                                }
                            )
                        if reconciliation_active and committed_count > index + 1:
                            progress["reconcile_element_count"] = committed_count
                        pending_in_progress = progress
                        yield StateMessage(
                            stream.name,
                            _checkpoint(
                                files_state,
                                parser_fingerprint,
                                progress,
                            ),
                        )
                        emitted = 0

                if resume_index > parsed_count:
                    raise ConfigurationError(
                        f"Document checkpoint next_index exceeds parsed elements for {path}"
                    )
                if committed_count > parsed_count:
                    if (
                        pending_in_progress is None
                        or pending_in_progress.get("relative_path") == relative_path
                    ):
                        tail_progress: dict[str, Any] = {
                            "relative_path": relative_path,
                            "sha256": fingerprint,
                            "next_index": parsed_count,
                            "reconcile_element_count": committed_count,
                            "parser_fingerprint": parser_fingerprint,
                        }
                        if manifest_key is not None and manifest_sha256 is not None:
                            tail_progress.update(
                                {
                                    "manifest_cache_key": manifest_key,
                                    "manifest_sha256": manifest_sha256,
                                }
                            )
                        if (
                            pending_in_progress is not None
                            and pending_in_progress.get("reconcile_element_count") is not None
                            and pending_in_progress.get("tombstone_next_index") is not None
                        ):
                            tail_progress["tombstone_next_index"] = pending_in_progress[
                                "tombstone_next_index"
                            ]
                        pending_in_progress = tail_progress
                    tombstone_start = parsed_count
                    if (
                        pending_in_progress is not None
                        and pending_in_progress.get("relative_path") == relative_path
                    ):
                        tombstone_start = max(
                            tombstone_start,
                            int(pending_in_progress.get("tombstone_next_index", 0)),
                        )
                    async for message in self._emit_tombstones(
                        stream.name,
                        relative_path,
                        committed_count,
                        files_state,
                        parser_fingerprint,
                        start=tombstone_start,
                        in_progress=pending_in_progress,
                    ):
                        yield message
                committed_file: dict[str, Any] = {
                    "sha256": fingerprint,
                    "element_count": parsed_count,
                    "parser_fingerprint": parser_fingerprint,
                }
                if manifest_key is not None and manifest_sha256 is not None:
                    committed_file.update(
                        {
                            "manifest_cache_key": manifest_key,
                            "manifest_sha256": manifest_sha256,
                        }
                    )
                files_state[relative_path] = committed_file
                if (
                    pending_in_progress is not None
                    and pending_in_progress.get("relative_path") == relative_path
                ):
                    pending_in_progress = None
                yield StateMessage(
                    stream.name,
                    _checkpoint(files_state, parser_fingerprint, pending_in_progress),
                )

        yield StateMessage(
            stream.name,
            _checkpoint(files_state, parser_fingerprint, pending_in_progress),
        )

    def _record_envelope(
        self,
        stream_name: str,
        original_path: Path,
        relative_path: str,
        fingerprint: str,
        stat: os.stat_result,
        index: int,
        element: _ParsedElement,
    ) -> Envelope:
        element_id = f"{relative_path}:{index}"
        event_time = datetime.fromtimestamp(stat.st_mtime, UTC)
        media_type = mimetypes.guess_type(original_path.name)[0] or "application/octet-stream"
        return Envelope(
            id=stable_record_id("local_documents", stream_name, element_id),
            source="local_documents",
            stream=stream_name,
            payload=element.payload,
            cursor=f"{fingerprint}:{index}",
            checksum=_payload_checksum(element.payload),
            event_time=event_time,
            metadata={
                **element.metadata,
                "element_id": element_id,
                "element_index": index,
                "relative_path": relative_path,
                "filename": original_path.name,
                "extension": original_path.suffix.lower(),
                "media_type": media_type,
                "size_bytes": stat.st_size,
                "modified_at": event_time.isoformat(),
            },
            provenance={
                "connector": "local_documents",
                "path": str(original_path.resolve()),
                "sha256": fingerprint,
            },
        )

    async def _emit_tombstones(
        self,
        stream_name: str,
        relative_path: str,
        count: int,
        files_state: Mapping[str, Mapping[str, Any]],
        parser_fingerprint: str,
        *,
        start: int = 0,
        in_progress: Mapping[str, Any] | None = None,
    ) -> AsyncIterator[SourceMessage]:
        emitted = 0
        for index in range(start, count):
            element_id = f"{relative_path}:{index}"
            yield RecordMessage(
                Envelope(
                    id=stable_record_id("local_documents", stream_name, element_id),
                    source="local_documents",
                    stream=stream_name,
                    payload=Tombstone("document element removed"),
                    operation=Operation.DELETE,
                    cursor=f"deleted:{relative_path}:{index}",
                    metadata={
                        "element_id": element_id,
                        "element_index": index,
                        "relative_path": relative_path,
                    },
                    provenance={"connector": "local_documents"},
                )
            )
            emitted += 1
            if emitted >= self.checkpoint_interval:
                checkpoint_progress = None if in_progress is None else dict(in_progress)
                if (
                    checkpoint_progress is not None
                    and checkpoint_progress.get("relative_path") == relative_path
                ):
                    checkpoint_progress["tombstone_next_index"] = index + 1
                yield StateMessage(
                    stream_name,
                    _checkpoint(files_state, parser_fingerprint, checkpoint_progress),
                )
                emitted = 0

    def _parser_fingerprint(self) -> str:
        legacy = {
            "parser_version": PARSER_VERSION,
            "text_chunk_chars": self.text_chunk_chars,
            "table_batch_rows": self.table_batch_rows,
        }
        if self.ocr_mode == "off" and self.table_mode == "off" and self.document_splitter is None:
            raw = json.dumps(legacy, sort_keys=True)
            return hashlib.sha256(raw.encode()).hexdigest()
        return canonical_fingerprint(
            {
                **legacy,
                "document_ai": {
                    "ocr_mode": self.ocr_mode,
                    "table_mode": self.table_mode,
                    "failure_mode": self.failure_mode,
                    "min_native_text_quality": self.min_native_text_quality,
                    "ocr_languages": self.ocr_languages,
                    "render_dpi": self.render_dpi,
                    "page_timeout_seconds": self.page_timeout_seconds,
                    "max_page_concurrency": self.max_page_concurrency,
                    "retain_extraction_artifacts": self.retain_extraction_artifacts,
                    "ocr_engine": _component_fingerprint(self.ocr_engine),
                    "page_renderer": _component_fingerprint(self.page_renderer),
                    "table_extractor": _component_fingerprint(self.table_extractor),
                    "document_splitter": _component_fingerprint(self.document_splitter),
                    "vision_extractor": _component_fingerprint(self.vision_extractor),
                    "vision_fallback": self.vision_fallback,
                    "vision_operation_version": (
                        "table-v1" if self.table_mode == "vision" or self.vision_fallback else None
                    ),
                    "vision_schema_version": (
                        "1" if self.table_mode == "vision" or self.vision_fallback else None
                    ),
                },
            }
        )

    def _requires_replay_manifest(self) -> bool:
        """Return whether any active component can change ordered extraction output."""
        active_components = [self.document_splitter]
        if self.ocr_mode != "off":
            active_components.extend((self.ocr_engine, self.page_renderer))
        if self.table_mode in {"auto", "local"}:
            active_components.extend((self.table_extractor, self.page_renderer))
        if self.table_mode == "vision" or self.vision_fallback:
            active_components.extend((self.vision_extractor, self.page_renderer))
        for component in active_components:
            descriptor = getattr(component, "descriptor", None)
            if descriptor is not None and (descriptor.external or not descriptor.deterministic):
                return True
        return False

    async def _parse_file_with_document_ai(
        self,
        path: Path,
    ) -> list[_ParsedElement]:
        """Run opt-in OCR/table/splitter processing against an immutable snapshot."""
        from ingestion_graph.document_ai import table_artifact_to_batches

        ocr_engine = self.ocr_engine or self._runtime_ocr_engine
        renderer = self.page_renderer or self._runtime_page_renderer
        table_extractor = self.table_extractor or self._runtime_table_extractor
        if self.ocr_mode != "off" and ocr_engine is None:
            from ingestion_graph.document_ai.tesseract import TesseractOcrEngine

            ocr_engine = TesseractOcrEngine()
            self._runtime_ocr_engine = ocr_engine
        if (
            (
                self.ocr_mode != "off"
                or self.table_mode in {"local", "auto", "vision"}
                or self.vision_fallback
            )
            and path.suffix.lower() == ".pdf"
            and renderer is None
        ):
            from ingestion_graph.document_ai.rendering import PdfiumPageRenderer

            renderer = PdfiumPageRenderer()
            self._runtime_page_renderer = renderer
        if self.table_mode in {"local", "auto"} and table_extractor is None:
            from ingestion_graph.document_ai.docling_adapter import DoclingTableExtractor

            table_extractor = DoclingTableExtractor()
            self._runtime_table_extractor = table_extractor
        component_locks: dict[int, asyncio.Lock] = {}

        def component_lock(component: object) -> asyncio.Lock:
            return component_locks.setdefault(id(component), asyncio.Lock())

        async def split_narrative(element: _ParsedElement) -> list[_ParsedElement]:
            if self.document_splitter is None or not isinstance(element.payload, DocumentElement):
                return [element]
            parent = element.payload
            parent_native_id = str(element.metadata.get("native_id", "element"))
            async with component_lock(self.document_splitter):
                chunks = await self.document_splitter.split(parent)
            split_elements: list[_ParsedElement] = []
            for index, chunk in enumerate(chunks):
                coordinates = (
                    parent.coordinates if chunk.coordinates is None else chunk.coordinates.to_dict()
                )
                metadata = {
                    **dict(chunk.metadata),
                    **dict(element.metadata),
                    "native_id": f"{parent_native_id}-chunk-{index}",
                    "splitter": _component_fingerprint(self.document_splitter),
                }
                if chunk.coordinates is None and coordinates is not None:
                    metadata["coordinate_precision"] = "parent"
                split_elements.append(
                    _ParsedElement(
                        DocumentElement(
                            chunk.text,
                            chunk.element_type,
                            parent.page_number if chunk.page_number is None else chunk.page_number,
                            chunk.parent_id or parent.parent_id or parent_native_id,
                            coordinates,
                        ),
                        metadata,
                    )
                )
            return split_elements

        async def process_image(
            image: bytes, page_number: int | None, native_text: str = ""
        ) -> list[_ParsedElement]:
            quality = evaluate_text_quality(native_text)
            should_ocr = self.ocr_mode == "always" or (
                self.ocr_mode == "auto" and quality.score < self.min_native_text_quality
            )
            results: list[_ParsedElement] = []
            text = native_text
            metadata: dict[str, Any] = {"page_number": page_number, "native_quality": quality.score}
            if should_ocr:
                if ocr_engine is None:
                    raise ConfigurationError("OCR requires an OCR engine")
                cache_key = canonical_fingerprint(
                    {
                        "image_sha256": hashlib.sha256(image).hexdigest(),
                        "engine": _component_fingerprint(ocr_engine),
                        "language": self.ocr_languages,
                    }
                )
                cached = (
                    await self.extraction_cache.get(cache_key)
                    if self.extraction_cache is not None and image
                    else None
                )
                if cached is not None:
                    try:
                        result = _deserialize_ocr_result(cached, cache_hit=True)
                    except (TypeError, ValueError, json.JSONDecodeError, UnicodeDecodeError):
                        if self.extraction_cache is None:
                            raise
                        await self.extraction_cache.delete(cache_key)
                        cached = None
                if cached is None:
                    async with component_lock(ocr_engine):
                        result = await ocr_engine.recognize(
                            image,
                            language="+".join(self.ocr_languages),
                        )
                    if self.extraction_cache is not None and image:
                        await self.extraction_cache.put(
                            cache_key,
                            _serialize_ocr_result(result),
                        )
                text = result.text
                metadata.update(
                    {
                        "extraction_mode": "ocr",
                        "ocr_confidence": result.confidence,
                        "ocr_engine": _component_fingerprint(ocr_engine),
                        "warnings": [item.to_dict() for item in result.warnings],
                    }
                )
            elif native_text:
                metadata["extraction_mode"] = "native"
            if text or page_number is not None:
                element = DocumentElement(text, "ocr_page" if should_ocr else "page", page_number)
                results.extend(
                    await split_narrative(
                        _ParsedElement(
                            element, {"native_id": f"page-{page_number or 1}", **metadata}
                        )
                    )
                )
            tables: Sequence[Any] = ()
            if table_extractor is not None and self.table_mode in {"local", "auto"}:
                async with component_lock(table_extractor):
                    tables = await table_extractor.extract(image, page_number=page_number)
            if (
                not tables
                and (self.vision_fallback or self.table_mode == "vision")
                and self.vision_extractor is not None
            ):
                if self.external_processing_policy is None:
                    raise ConfigurationError("External vision processing was not authorized")
                async with component_lock(self.external_processing_policy):
                    authorized = await self.external_processing_policy.authorize(
                        purpose="table", page_number=page_number
                    )
                if not authorized:
                    raise ConfigurationError("External vision processing was not authorized")
                from ingestion_graph.document_ai import (
                    VISION_TABLE_RESPONSE_SCHEMA,
                    validate_vision_table_response,
                )

                candidate_tables: Sequence[Any] | None = None
                for _attempt in range(2):
                    async with component_lock(self.vision_extractor):
                        candidate = await self.vision_extractor.extract(
                            image,
                            schema=VISION_TABLE_RESPONSE_SCHEMA,
                        )
                    if not isinstance(candidate, Mapping):
                        continue
                    try:
                        candidate_tables = validate_vision_table_response(candidate)
                    except ConfigurationError:
                        continue
                    else:
                        break
                if candidate_tables is None:
                    if self.failure_mode == "strict":
                        raise ConfigurationError(
                            "Vision extractor returned invalid table artifacts"
                        )
                    return results
                tables = tuple(candidate_tables)
            from ingestion_graph.document_ai import TableArtifact

            for table_index, artifact in enumerate(tables):
                if isinstance(artifact, Mapping):
                    artifact = TableArtifact.from_dict(artifact)
                for batch_index, batch in enumerate(
                    table_artifact_to_batches(artifact, batch_rows=self.table_batch_rows)
                ):
                    table_metadata: dict[str, Any] = {
                        "native_id": (f"table-{page_number or 1}-{table_index}-{batch_index}"),
                        "page_number": page_number,
                        "table_index": table_index,
                        "batch_index": batch_index,
                        "table_id": artifact.table_id,
                        "table_confidence": artifact.confidence,
                        "table_warnings": [warning.to_dict() for warning in artifact.warnings],
                    }
                    if self.retain_extraction_artifacts:
                        table_metadata["table_artifact"] = artifact.to_dict()
                    results.append(
                        _ParsedElement(
                            batch,
                            table_metadata,
                        )
                    )
            return results

        if path.suffix.lower() in OCR_IMAGE_EXTENSIONS:
            from ingestion_graph.document_ai import validate_image_payload

            async def process_direct_image() -> list[_ParsedElement]:
                image = await asyncio.to_thread(path.read_bytes)
                await asyncio.to_thread(
                    validate_image_payload,
                    image,
                    extension=path.suffix.lower(),
                    max_frames=(1 if self.table_mode == "vision" or self.vision_fallback else 256),
                )
                return await process_image(image, 1)

            return await asyncio.wait_for(
                process_direct_image(),
                timeout=self.page_timeout_seconds,
            )
        if path.suffix.lower() != ".pdf":
            parsed = list(
                _parse_file(
                    path,
                    text_chunk_chars=self.text_chunk_chars,
                    table_batch_rows=self.table_batch_rows,
                )
            )
            split: list[_ParsedElement] = []
            for element in parsed:
                split.extend(await split_narrative(element))
            return split
        if PdfReader is None:
            raise ConfigurationError(
                "PDF support requires: pip install 'ingestion-graph[documents]'"
            )
        reader = PdfReader(str(path))
        pdf_bytes = await asyncio.to_thread(path.read_bytes)
        semaphore = asyncio.Semaphore(self.max_page_concurrency)

        async def process_pdf_page(page_number: int, native_text: str) -> list[_ParsedElement]:
            async with semaphore:

                async def run_page() -> list[_ParsedElement]:
                    quality = evaluate_text_quality(native_text)
                    needs_render = (
                        self.ocr_mode == "always"
                        or (
                            self.ocr_mode == "auto" and quality.score < self.min_native_text_quality
                        )
                        or self.table_mode in {"local", "auto"}
                        or self.table_mode == "vision"
                        or self.vision_fallback
                    )
                    if not needs_render:
                        return await process_image(b"", page_number, native_text)
                    if renderer is None:
                        raise ConfigurationError(
                            "PDF OCR/table processing requires a page renderer"
                        )
                    async with component_lock(renderer):
                        image = await renderer.render(
                            pdf_bytes,
                            page_number=page_number,
                            dpi=self.render_dpi,
                        )
                    return await process_image(image, page_number, native_text)

                return await asyncio.wait_for(
                    run_page(),
                    timeout=self.page_timeout_seconds,
                )

        elements: list[_ParsedElement] = []
        for start in range(0, len(reader.pages), self.max_page_concurrency):
            stop = min(start + self.max_page_concurrency, len(reader.pages))
            window = [
                (index + 1, reader.pages[index].extract_text() or "")
                for index in range(start, stop)
            ]
            tasks = [
                asyncio.create_task(process_pdf_page(page_number, native_text))
                for page_number, native_text in window
            ]
            try:
                page_elements = await asyncio.gather(*tasks)
            except BaseException:
                for task in tasks:
                    if not task.done():
                        task.cancel()
                await asyncio.gather(*tasks, return_exceptions=True)
                raise
            elements.extend(element for page in page_elements for element in page)
        return elements

    async def close(self) -> None:
        if self._closed:
            return
        self._closed = True
        components = (
            self.ocr_engine,
            self.page_renderer,
            self.table_extractor,
            self.document_splitter,
            self.vision_extractor,
            self.extraction_cache,
            self._runtime_ocr_engine,
            self._runtime_page_renderer,
            self._runtime_table_extractor,
        )
        closed: set[int] = set()
        first_error: BaseException | None = None
        for component in components:
            if component is None or id(component) in closed:
                continue
            closed.add(id(component))
            close = getattr(component, "close", None)
            if callable(close):
                try:
                    await close()
                except BaseException as exc:
                    if first_error is None:
                        first_error = exc
        if first_error is not None:
            raise first_error

    def _discover_files(self, only_root: Path | None = None) -> list[Path]:
        files: list[Path] = []
        for root in (only_root,) if only_root is not None else self.paths:
            if not self.follow_symlinks and root.is_symlink():
                raise ConfigurationError(f"Configured document root must not be a symlink: {root}")
            if root.is_file():
                candidates: Iterable[Path] = (root,)
            elif root.is_dir():
                candidates = self._walk_directory(root)
            else:
                continue
            for candidate in candidates:
                if not candidate.is_file() or candidate.suffix.lower() not in self.extensions:
                    continue
                if not self.follow_symlinks and candidate.is_symlink():
                    continue
                if not self.follow_symlinks:
                    try:
                        candidate.resolve(strict=True).relative_to(root.resolve(strict=True))
                    except (OSError, ValueError):
                        continue
                if not self.include_hidden and _is_hidden(candidate, root):
                    continue
                files.append(candidate)
        return sorted(set(files), key=lambda item: str(item).casefold())

    def _walk_directory(self, root: Path) -> Iterable[Path]:
        visited: set[tuple[int, int]] = set()
        for current, directory_names, file_names in os.walk(root, followlinks=self.follow_symlinks):
            current_path = Path(current)
            if self.follow_symlinks:
                try:
                    current_stat = current_path.stat()
                except OSError:
                    directory_names.clear()
                    continue
                identity = (current_stat.st_dev, current_stat.st_ino)
                if identity in visited:
                    directory_names.clear()
                    continue
                visited.add(identity)
            if not self.include_hidden:
                directory_names[:] = [name for name in directory_names if not name.startswith(".")]
            if not self.follow_symlinks:
                directory_names[:] = [
                    name for name in directory_names if not (current_path / name).is_symlink()
                ]
            if not self.recursive:
                directory_names.clear()
            yield from (current_path / name for name in file_names)

    def _name_streams(self) -> dict[str, Path]:
        streams: dict[str, Path] = {}
        for index, root in enumerate(self.paths):
            if self.stream_names:
                name = self.stream_names[index]
            else:
                base = root.stem if root.suffix.lower() in self.extensions else root.name
                path_hash = hashlib.sha256(str(root.resolve()).encode()).hexdigest()[:10]
                name = f"{base or 'documents'}#{path_hash}"
            streams[name] = root
        return streams


def _normalize_extension(value: str) -> str:
    normalized = str(value).strip().lower()
    if not normalized:
        raise ConfigurationError("Document extension must not be empty")
    return normalized if normalized.startswith(".") else f".{normalized}"


def _relative_path(path: Path, root: Path) -> str:
    return path.name if root.is_file() else path.relative_to(root).as_posix()


def _is_hidden(path: Path, root: Path) -> bool:
    boundary = root.parent if root.is_file() else root
    try:
        relative = path.relative_to(boundary)
    except ValueError:
        relative = path
    return any(part.startswith(".") for part in relative.parts)


def _trusted_root(root: Path) -> _TrustedRoot | None:
    """Capture a stable root identity before discovery without following a root link."""
    try:
        before = root.lstat()
    except FileNotFoundError:
        return None
    if _is_reparse(root):
        raise ConfigurationError(f"Configured document root must not be a symlink: {root}")
    resolved = root.resolve(strict=True)
    after = root.lstat()
    if _stat_identity(before) != _stat_identity(after) or _is_reparse(root):
        raise ConfigurationError(f"Configured document root changed during discovery: {root}")
    return _TrustedRoot(resolved, _stat_identity(after), root.is_file())


def _stat_identity(value: os.stat_result) -> tuple[int, int]:
    return value.st_dev, value.st_ino


def _is_reparse(path: Path) -> bool:
    try:
        value = path.lstat()
    except OSError:
        return False
    attributes = getattr(value, "st_file_attributes", 0)
    reparse_flag = getattr(stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400)
    return path.is_symlink() or bool(attributes & reparse_flag)


def _assert_contained(candidate: Path, root: Path, trusted_root: _TrustedRoot) -> Path:
    resolved = candidate.resolve(strict=True)
    if trusted_root.is_file:
        if resolved != trusted_root.path:
            raise ConfigurationError(f"Document path escaped its configured root: {candidate}")
    else:
        try:
            resolved.relative_to(trusted_root.path)
        except ValueError as exc:
            raise ConfigurationError(
                f"Document path escaped its configured root: {candidate}"
            ) from exc

        current = root
        try:
            relative_parts = candidate.relative_to(root).parts
        except ValueError as exc:
            raise ConfigurationError(
                f"Document path escaped its configured root: {candidate}"
            ) from exc
        if _is_reparse(root):
            raise ConfigurationError(f"Configured document root must not be a symlink: {root}")
        for part in relative_parts:
            current = current / part
            if _is_reparse(current):
                raise ConfigurationError(f"Document path contains a symlink: {candidate}")
    return resolved


def _open_source(
    path: Path,
    *,
    root: Path,
    trusted_root: _TrustedRoot | None,
    follow_symlinks: bool,
) -> int:
    file_flags = os.O_RDONLY | getattr(os, "O_BINARY", 0) | getattr(os, "O_NONBLOCK", 0)
    if follow_symlinks:
        return os.open(path, file_flags)
    if trusted_root is None:
        raise ConfigurationError(f"Configured document root disappeared: {root}")

    no_follow = getattr(os, "O_NOFOLLOW", 0)
    if os.name == "nt" or os.open not in os.supports_dir_fd:
        return os.open(path, file_flags | no_follow)
    if trusted_root.is_file:
        descriptor = os.open(root, file_flags | no_follow)
        if _stat_identity(os.fstat(descriptor)) != trusted_root.identity:
            os.close(descriptor)
            raise ConfigurationError(f"Configured document root changed while opening: {root}")
        return descriptor

    directory_flags = os.O_RDONLY | no_follow | getattr(os, "O_DIRECTORY", 0)
    directory_descriptor = os.open(root, directory_flags)
    try:
        if _stat_identity(os.fstat(directory_descriptor)) != trusted_root.identity:
            raise ConfigurationError(f"Configured document root changed while opening: {root}")
        parts = path.relative_to(root).parts
        if not parts:
            raise ConfigurationError(f"Document path is not a file below its root: {path}")
        for part in parts[:-1]:
            next_descriptor = os.open(
                part,
                directory_flags,
                dir_fd=directory_descriptor,
            )
            os.close(directory_descriptor)
            directory_descriptor = next_descriptor
        return os.open(
            parts[-1],
            file_flags | no_follow,
            dir_fd=directory_descriptor,
        )
    finally:
        os.close(directory_descriptor)


@contextmanager
def _snapshot_file(
    path: Path,
    *,
    root: Path,
    trusted_root: _TrustedRoot | None,
    follow_symlinks: bool,
    max_bytes: int,
) -> Iterator[_Snapshot]:
    source_descriptor: int | None = None
    temp_descriptor: int | None = None
    temp_name: str | None = None
    digest = hashlib.sha256()
    try:
        expected_path = (
            _assert_contained(path, root, trusted_root)
            if not follow_symlinks and trusted_root is not None
            else path.resolve(strict=True)
        )
        source_descriptor = _open_source(
            path,
            root=root,
            trusted_root=trusted_root,
            follow_symlinks=follow_symlinks,
        )
        temp_descriptor, temp_name = tempfile.mkstemp(
            prefix="ingestion-document-", suffix=path.suffix
        )
        with (
            os.fdopen(source_descriptor, "rb") as source,
            os.fdopen(temp_descriptor, "wb") as target,
        ):
            source_descriptor = None
            temp_descriptor = None
            source_stat = os.fstat(source.fileno())
            if not stat.S_ISREG(source_stat.st_mode):
                raise ConfigurationError(f"Document path is not a regular file: {path}")
            if source_stat.st_size > max_bytes:
                raise ConfigurationError(f"Document file exceeds the configured size limit: {path}")
            if not follow_symlinks:
                if trusted_root is None:
                    raise ConfigurationError(f"Configured document root disappeared: {root}")
                current_path = _assert_contained(path, root, trusted_root)
                try:
                    path_stat = path.lstat()
                except OSError as exc:
                    raise ConfigurationError(
                        f"Document path changed while opening: {path}"
                    ) from exc
                if current_path != expected_path or _stat_identity(path_stat) != _stat_identity(
                    source_stat
                ):
                    raise ConfigurationError(f"Document path changed while opening: {path}")
            copied = 0
            while block := source.read(1024 * 1024):
                copied += len(block)
                if copied > max_bytes:
                    raise ConfigurationError(
                        f"Document file exceeds the configured size limit: {path}"
                    )
                digest.update(block)
                target.write(block)
            target.flush()
        yield _Snapshot(Path(temp_name), digest.hexdigest(), source_stat)
    finally:
        if source_descriptor is not None:
            with suppress(OSError):
                os.close(source_descriptor)
        if temp_descriptor is not None:
            with suppress(OSError):
                os.close(temp_descriptor)
        if temp_name is not None:
            Path(temp_name).unlink(missing_ok=True)


def _validate_state(state: Mapping[str, Any] | None) -> dict[str, Any]:
    current = dict(state or {})
    raw_files = current.get("files", {})
    if not isinstance(raw_files, Mapping):
        raise ConfigurationError("Document checkpoint files must be an object")
    files: dict[str, dict[str, Any]] = {}
    for relative_path, raw_file in raw_files.items():
        if not isinstance(relative_path, str) or not isinstance(raw_file, Mapping):
            raise ConfigurationError("Document checkpoint file entries are invalid")
        sha256 = raw_file.get("sha256")
        element_count = raw_file.get("element_count")
        if not isinstance(sha256, str):
            raise ConfigurationError("Document checkpoint file sha256 must be a string")
        if (
            isinstance(element_count, bool)
            or not isinstance(element_count, int)
            or element_count < 0
        ):
            raise ConfigurationError("Document checkpoint file element_count must be non-negative")
    parser_fingerprint = current.get("parser_fingerprint")
    if parser_fingerprint is not None and not isinstance(parser_fingerprint, str):
        raise ConfigurationError("Document checkpoint parser_fingerprint must be a string")
    for relative_path, raw_file in raw_files.items():
        file_parser = raw_file.get("parser_fingerprint", parser_fingerprint)
        if file_parser is not None and not isinstance(file_parser, str):
            raise ConfigurationError("Document checkpoint file parser_fingerprint must be a string")
        files[relative_path] = {
            "sha256": raw_file["sha256"],
            "element_count": raw_file["element_count"],
            "parser_fingerprint": file_parser,
        }
        for key in ("manifest_cache_key", "manifest_sha256"):
            item = raw_file.get(key)
            if item is not None:
                if not isinstance(item, str) or not item:
                    raise ConfigurationError(f"Document checkpoint file {key} must be a string")
                files[relative_path][key] = item
    in_progress = current.get("in_progress")
    if in_progress is not None:
        if not isinstance(in_progress, Mapping):
            raise ConfigurationError("Document checkpoint in_progress must be an object")
        relative_path = in_progress.get("relative_path")
        sha256 = in_progress.get("sha256")
        next_index = in_progress.get("next_index")
        progress_parser = in_progress.get("parser_fingerprint", parser_fingerprint)
        if not isinstance(relative_path, str) or not isinstance(sha256, str):
            raise ConfigurationError("Document checkpoint in_progress identity is invalid")
        if isinstance(next_index, bool) or not isinstance(next_index, int) or next_index < 0:
            raise ConfigurationError(
                "Document checkpoint in_progress next_index must be non-negative"
            )
        if progress_parser is not None and not isinstance(progress_parser, str):
            raise ConfigurationError(
                "Document checkpoint in_progress parser_fingerprint must be a string"
            )
        in_progress = {
            "relative_path": relative_path,
            "sha256": sha256,
            "next_index": next_index,
            "parser_fingerprint": progress_parser,
        }
        for key in ("manifest_cache_key", "manifest_sha256"):
            item = current["in_progress"].get(key)
            if item is not None:
                if not isinstance(item, str) or not item:
                    raise ConfigurationError(
                        f"Document checkpoint in_progress {key} must be a string"
                    )
                in_progress[key] = item
        reconcile_element_count = current["in_progress"].get("reconcile_element_count")
        if reconcile_element_count is not None:
            if (
                isinstance(reconcile_element_count, bool)
                or not isinstance(reconcile_element_count, int)
                or reconcile_element_count < next_index
            ):
                raise ConfigurationError(
                    "Document checkpoint in_progress reconcile_element_count is invalid"
                )
            in_progress["reconcile_element_count"] = reconcile_element_count
        tombstone_next_index = current["in_progress"].get("tombstone_next_index")
        if tombstone_next_index is not None:
            tombstone_bound = (
                reconcile_element_count
                if isinstance(reconcile_element_count, int)
                and not isinstance(reconcile_element_count, bool)
                else next_index
            )
            if (
                isinstance(tombstone_next_index, bool)
                or not isinstance(tombstone_next_index, int)
                or tombstone_next_index < 0
                or tombstone_next_index > tombstone_bound
            ):
                raise ConfigurationError(
                    "Document checkpoint in_progress tombstone_next_index is invalid"
                )
            in_progress["tombstone_next_index"] = tombstone_next_index
    return {
        "files": files,
        "parser_fingerprint": parser_fingerprint,
        "in_progress": in_progress,
    }


def _checkpoint(
    files: Mapping[str, Mapping[str, Any]],
    parser_fingerprint: str,
    in_progress: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    value: dict[str, Any] = {
        "parser_fingerprint": parser_fingerprint,
        "files": {key: dict(item) for key, item in files.items()},
    }
    if in_progress is not None:
        value["in_progress"] = dict(in_progress)
    return value


def _missing_dependency(files: Sequence[Path]) -> str | None:
    extensions = {path.suffix.lower() for path in files}
    if ".pdf" in extensions and PdfReader is None:
        return "PDF support requires: pip install 'ingestion-graph[documents]'"
    if ".docx" in extensions and WordDocument is None:
        return "Word support requires: pip install 'ingestion-graph[documents]'"
    if ".xlsx" in extensions and load_workbook is None:
        return "Excel support requires: pip install 'ingestion-graph[documents]'"
    return None


def _validate_archive(path: Path, *, max_uncompressed_bytes: int) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            total = sum(item.file_size for item in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise ConfigurationError(f"Document archive is invalid: {path.name}") from exc
    if total > max_uncompressed_bytes:
        raise ConfigurationError(
            f"Document archive exceeds the uncompressed size limit: {path.name}"
        )


def _parse_file(
    path: Path,
    *,
    text_chunk_chars: int,
    table_batch_rows: int,
) -> Iterable[_ParsedElement]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path, text_chunk_chars, table_batch_rows)
    if extension == ".xlsx":
        return _parse_xlsx(path, table_batch_rows)
    if extension == ".csv":
        return _parse_csv(path, table_batch_rows)
    if extension == ".eml":
        return _parse_email(path, text_chunk_chars)
    text = path.read_text(encoding="utf-8", errors="replace")
    if extension in {".htm", ".html"}:
        parser = _TextHTMLParser()
        parser.feed(text)
        text = parser.text
    return _document_chunks(text, text_chunk_chars, element_type="text")


def _parse_pdf(path: Path) -> Iterable[_ParsedElement]:
    if PdfReader is None:
        raise ConfigurationError("PDF support requires: pip install 'ingestion-graph[documents]'")
    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        yield _ParsedElement(
            DocumentElement(
                text=page.extract_text() or "",
                element_type="page",
                page_number=index,
            ),
            {"native_id": f"page-{index}", "page_number": index},
        )


def _parse_docx(path: Path, chunk_chars: int, batch_rows: int) -> list[_ParsedElement]:
    if WordDocument is None:
        raise ConfigurationError("Word support requires: pip install 'ingestion-graph[documents]'")
    document = WordDocument(str(path))
    paragraphs = "\n\n".join(item.text for item in document.paragraphs if item.text.strip())
    elements = _document_chunks(paragraphs, chunk_chars, element_type="paragraphs")
    for table_index, table in enumerate(document.tables, start=1):
        values = [[cell.text for cell in row.cells] for row in table.rows]
        if not values:
            continue
        columns = _column_names(values[0])
        rows = [dict(zip(columns, row, strict=False)) for row in values[1:]]
        for batch_index, offset in enumerate(range(0, len(rows), batch_rows)):
            batch = rows[offset : offset + batch_rows]
            elements.append(
                _ParsedElement(
                    TableBatch(columns, batch),
                    {
                        "native_id": f"table-{table_index}-batch-{batch_index}",
                        "table_index": table_index,
                        "batch_index": batch_index,
                        "row_count": len(batch),
                    },
                )
            )
    return elements


def _parse_xlsx(path: Path, batch_rows: int) -> Iterable[_ParsedElement]:
    if load_workbook is None:
        raise ConfigurationError("Excel support requires: pip install 'ingestion-graph[documents]'")
    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    try:
        for worksheet in workbook.worksheets:
            iterator = worksheet.iter_rows(values_only=True)
            first = next(iterator, None)
            if first is None:
                continue
            columns = _column_names(tuple(_json_cell(value) for value in first))
            rows: list[Mapping[str, Any]] = []
            batch_index = 0
            for values in iterator:
                rows.append(
                    dict(
                        zip(
                            columns,
                            (_json_cell(value) for value in values),
                            strict=False,
                        )
                    )
                )
                if len(rows) >= batch_rows:
                    yield _table_element(worksheet.title, batch_index, columns, rows)
                    rows = []
                    batch_index += 1
            if rows:
                yield _table_element(worksheet.title, batch_index, columns, rows)
    finally:
        workbook.close()


def _parse_csv(path: Path, batch_rows: int) -> Iterable[_ParsedElement]:
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        header = next(reader, ())
        columns = _column_names(header)
        rows: list[Mapping[str, Any]] = []
        batch_index = 0
        for values in reader:
            padded = [*values, *([None] * max(0, len(columns) - len(values)))]
            rows.append(dict(zip(columns, padded, strict=False)))
            if len(rows) >= batch_rows:
                yield _table_element("csv", batch_index, columns, rows)
                rows = []
                batch_index += 1
        if rows:
            yield _table_element("csv", batch_index, columns, rows)


def _parse_email(path: Path, chunk_chars: int) -> list[_ParsedElement]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    body = message.get_body(preferencelist=("plain", "html"))
    text = body.get_content() if body is not None else ""
    if body is not None and body.get_content_type() == "text/html":
        parser = _TextHTMLParser()
        parser.feed(str(text))
        text = parser.text
    attachments = [
        {
            "filename": part.get_filename(),
            "content_type": part.get_content_type(),
            "size_bytes": len(part.get_payload(decode=True) or b""),
        }
        for part in message.iter_attachments()
    ]
    common = {
        "subject": str(message.get("subject", "")),
        "from": str(message.get("from", "")),
        "to": str(message.get("to", "")),
        "date": str(message.get("date", "")),
        "message_id": str(message.get("message-id", "")),
        "attachments": attachments,
    }
    chunks = _document_chunks(str(text), chunk_chars, element_type="email_body")
    return [_ParsedElement(item.payload, {**common, **item.metadata}) for item in chunks] or [
        _ParsedElement(DocumentElement("", "email_body"), {"native_id": "body-0", **common})
    ]


def _document_chunks(text: str, limit: int, *, element_type: str) -> list[_ParsedElement]:
    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        remaining_length = len(normalized) - start
        if remaining_length <= limit:
            chunks.append(normalized[start:])
            break
        end = start + limit
        split_at = normalized.rfind("\n", start, end + 1)
        if split_at < start + limit // 2:
            split_at = normalized.rfind(" ", start, end + 1)
        if split_at < start + limit // 2:
            split_at = end
        chunks.append(normalized[start:split_at].strip())
        start = split_at
        while start < len(normalized) and normalized[start].isspace():
            start += 1
    return [
        _ParsedElement(
            DocumentElement(chunk, element_type=element_type),
            {"native_id": f"{element_type}-{index}", "chunk_index": index},
        )
        for index, chunk in enumerate(chunks)
    ]


def _column_names(values: Sequence[Any]) -> tuple[str, ...]:
    columns: list[str] = []
    seen: dict[str, int] = {}
    for index, value in enumerate(values, start=1):
        base = str(value).strip() if value is not None else ""
        base = base or f"column_{index}"
        count = seen.get(base, 0) + 1
        seen[base] = count
        columns.append(base if count == 1 else f"{base}_{count}")
    return tuple(columns)


def _json_cell(value: Any) -> Any:
    if value is None or isinstance(value, (str, bool, int, float)):
        return value
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, (date, time)):
        return value.isoformat()
    if isinstance(value, timedelta):
        return value.total_seconds()
    return str(value)


def _table_element(
    sheet: str,
    batch_index: int,
    columns: Sequence[str],
    rows: Sequence[Mapping[str, Any]],
) -> _ParsedElement:
    return _ParsedElement(
        TableBatch(tuple(columns), tuple(dict(row) for row in rows)),
        {
            "native_id": f"{sheet}-batch-{batch_index}",
            "sheet": sheet,
            "batch_index": batch_index,
            "row_count": len(rows),
        },
    )


def _serialize_extraction_manifest(elements: Sequence[_ParsedElement]) -> bytes:
    items: list[dict[str, Any]] = []
    for element in elements:
        payload = Envelope(
            id="manifest",
            source="local_documents",
            stream="manifest",
            payload=element.payload,
        ).to_dict()["payload"]
        items.append({"payload": payload, "metadata": dict(element.metadata)})
    return json.dumps(
        {"schema_version": "1", "elements": items},
        sort_keys=True,
        separators=(",", ":"),
        ensure_ascii=False,
    ).encode("utf-8")


def _deserialize_extraction_manifest(value: bytes) -> list[_ParsedElement]:
    document = json.loads(value.decode("utf-8"))
    if not isinstance(document, Mapping) or document.get("schema_version") != "1":
        raise ValueError("Unsupported extraction manifest schema")
    raw_elements = document.get("elements")
    if not isinstance(raw_elements, list):
        raise ValueError("Extraction manifest elements must be an array")
    elements: list[_ParsedElement] = []
    for raw in raw_elements:
        if not isinstance(raw, Mapping) or not isinstance(raw.get("payload"), Mapping):
            raise ValueError("Extraction manifest element is invalid")
        metadata = raw.get("metadata", {})
        if not isinstance(metadata, Mapping):
            raise ValueError("Extraction manifest metadata must be an object")
        envelope = Envelope.from_dict(
            {
                "id": "manifest",
                "source": "local_documents",
                "stream": "manifest",
                "payload": raw["payload"],
            }
        )
        elements.append(_ParsedElement(envelope.payload, dict(metadata)))
    return elements


def _serialize_ocr_result(result: Any) -> bytes:
    value = {
        "schema_version": "1",
        "text": result.text,
        "confidence": result.confidence,
        "tokens": [
            {
                "text": token.text,
                "confidence": token.confidence,
                "coordinates": None if token.coordinates is None else token.coordinates.to_dict(),
            }
            for token in result.tokens
        ],
        "warnings": [warning.to_dict() for warning in result.warnings],
        "usage": dict(result.usage),
    }
    return json.dumps(value, sort_keys=True, separators=(",", ":")).encode("utf-8")


def _deserialize_ocr_result(value: bytes, *, cache_hit: bool) -> Any:
    from ingestion_graph.document_ai import (
        BoundingBox,
        ExtractionWarning,
        OcrResult,
        OcrToken,
    )

    raw = json.loads(value.decode("utf-8"))
    if not isinstance(raw, Mapping) or raw.get("schema_version") != "1":
        raise ValueError("Unsupported OCR cache schema")
    raw_tokens = raw.get("tokens", [])
    raw_warnings = raw.get("warnings", [])
    if not isinstance(raw_tokens, list) or not isinstance(raw_warnings, list):
        raise ValueError("OCR cache entry is invalid")
    tokens = []
    for item in raw_tokens:
        if not isinstance(item, Mapping):
            raise ValueError("OCR cache token is invalid")
        coordinates = item.get("coordinates")
        tokens.append(
            OcrToken(
                text=str(item.get("text", "")),
                confidence=item.get("confidence"),
                coordinates=None
                if not isinstance(coordinates, Mapping)
                else BoundingBox(**coordinates),
            )
        )
    warnings = []
    for item in raw_warnings:
        if not isinstance(item, Mapping):
            raise ValueError("OCR cache warning is invalid")
        warnings.append(
            ExtractionWarning(
                code=str(item.get("code", "ocr_warning")),
                message=str(item.get("message", "OCR warning")),
                page_number=item.get("page_number")
                if isinstance(item.get("page_number"), int)
                else None,
            )
        )
    usage = raw.get("usage", {})
    if not isinstance(usage, Mapping):
        raise ValueError("OCR cache usage is invalid")
    return OcrResult(
        text=str(raw.get("text", "")),
        tokens=tuple(tokens),
        confidence=raw.get("confidence"),
        warnings=tuple(warnings),
        usage={**dict(usage), "cache_hit": cache_hit},
    )


def _payload_checksum(payload: Payload) -> str:
    if isinstance(payload, DocumentElement):
        value: Any = {
            "text": payload.text,
            "element_type": payload.element_type,
            "page_number": payload.page_number,
        }
    elif isinstance(payload, TableBatch):
        value = {"columns": list(payload.columns), "rows": list(payload.rows)}
    else:
        value = repr(payload)
    encoded = json.dumps(value, sort_keys=True, default=str, ensure_ascii=False).encode()
    return hashlib.sha256(encoded).hexdigest()


def _component_fingerprint(component: Any) -> Any:
    if component is None:
        return None
    descriptor = getattr(component, "descriptor", None)
    if descriptor is not None and hasattr(descriptor, "to_dict"):
        return descriptor.to_dict()
    if isinstance(component, (str, int, float, bool)):
        return component
    return {"type": type(component).__module__ + "." + type(component).__qualname__}


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style"}:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style"} and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth and data.strip():
            self._parts.append(data.strip())

    @property
    def text(self) -> str:
        return "\n".join(self._parts)
