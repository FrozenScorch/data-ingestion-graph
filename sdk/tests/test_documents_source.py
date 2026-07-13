from __future__ import annotations

import hashlib
from datetime import date, datetime, time, timedelta
from email.message import EmailMessage
from pathlib import Path

import pytest

from ingestion_graph import ConnectorSpec
from ingestion_graph import LocalDocumentsSource as PublicLocalDocumentsSource
from ingestion_graph.destinations import JsonlDestination, SQLiteCollection
from ingestion_graph.errors import ConfigurationError
from ingestion_graph.messages import RecordMessage, StateMessage
from ingestion_graph.models import DocumentElement, Operation, TableBatch
from ingestion_graph.pipeline import Pipeline
from ingestion_graph.query import QueryRequest
from ingestion_graph.sources import LocalDocumentsSource, documents
from ingestion_graph.state import MemoryStateStore


async def collect(source: LocalDocumentsSource, stream, state=None):
    return [message async for message in source.read(stream, state)]


def records(messages):
    return [message for message in messages if isinstance(message, RecordMessage)]


def states(messages):
    return [message for message in messages if isinstance(message, StateMessage)]


@pytest.mark.asyncio
async def test_discovers_one_stable_stream_per_root_and_filters_files(tmp_path: Path):
    (tmp_path / "notes.txt").write_text("hello", encoding="utf-8")
    (tmp_path / "data.csv").write_text("id,name\n1,Ada\n", encoding="utf-8")
    (tmp_path / "ignore.bin").write_bytes(b"binary")
    (tmp_path / ".hidden.md").write_text("secret", encoding="utf-8")
    nested = tmp_path / "nested"
    nested.mkdir()
    (nested / "readme.md").write_text("nested", encoding="utf-8")

    source = LocalDocumentsSource(tmp_path, stream_names=["personal-documents"])
    assert PublicLocalDocumentsSource is LocalDocumentsSource
    assert isinstance(source.spec(), ConnectorSpec)
    assert (await source.check()).ok
    streams = await source.discover()
    assert [stream.name for stream in streams] == ["personal-documents"]

    messages = await collect(source, streams[0])
    filenames = {message.envelope.metadata["filename"] for message in records(messages)}
    assert filenames == {"data.csv", "readme.md", "notes.txt"}
    assert source.spec().capabilities.deletes is True


@pytest.mark.asyncio
async def test_text_resume_and_parser_config_are_checkpoint_safe(tmp_path: Path):
    path = tmp_path / "long.txt"
    path.write_text(("first section " * 30) + "\n" + ("second section " * 30), encoding="utf-8")
    source = LocalDocumentsSource(
        path,
        stream_names=["notes"],
        checkpoint_interval=1,
        text_chunk_chars=256,
    )
    stream = (await source.discover())[0]

    first = await collect(source, stream)
    first_records = records(first)
    first_states = states(first)
    partial = next(state.state for state in first_states if "in_progress" in state.state)
    final = first_states[-1].state
    assert len(first_records) >= 2

    resumed = await collect(source, stream, partial)
    assert [item.envelope.id for item in records(resumed)] == [
        item.envelope.id for item in first_records[1:]
    ]
    assert not records(await collect(source, stream, final))

    rechunked = LocalDocumentsSource(path, stream_names=["notes"], text_chunk_chars=12_000)
    rechunked_stream = (await rechunked.discover())[0]
    changed = await collect(rechunked, rechunked_stream, final)
    assert any(item.envelope.operation is Operation.UPSERT for item in records(changed))
    assert any(item.envelope.operation is Operation.DELETE for item in records(changed))
    assert states(changed)[-1].state["parser_fingerprint"] != final["parser_fingerprint"]


@pytest.mark.asyncio
async def test_snapshot_hash_always_matches_bytes_that_are_parsed(tmp_path: Path, monkeypatch):
    path = tmp_path / "race.txt"
    path.write_text("old bytes", encoding="utf-8")
    real_snapshot = documents._snapshot_file

    def mutate_before_snapshot(candidate: Path, **kwargs):
        path.write_text("new bytes", encoding="utf-8")
        return real_snapshot(candidate, **kwargs)

    monkeypatch.setattr(documents, "_snapshot_file", mutate_before_snapshot)
    source = LocalDocumentsSource(path, stream_names=["race"])
    stream = (await source.discover())[0]
    messages = await collect(source, stream)

    assert records(messages)[0].envelope.payload.text == "new bytes"
    expected = hashlib.sha256(b"new bytes").hexdigest()
    assert states(messages)[-1].state["files"]["race.txt"]["sha256"] == expected


@pytest.mark.asyncio
async def test_interrupted_parser_migration_reprocesses_every_file(tmp_path: Path):
    for name in ("a.txt", "b.txt"):
        (tmp_path / name).write_text("many words " * 300, encoding="utf-8")
    original = LocalDocumentsSource(
        tmp_path,
        stream_names=["documents"],
        text_chunk_chars=256,
    )
    stream = (await original.discover())[0]
    old_state = states(await collect(original, stream))[-1].state

    migrated = LocalDocumentsSource(
        tmp_path,
        stream_names=["documents"],
        text_chunk_chars=12_000,
    )
    iterator = migrated.read(stream, old_state)
    interrupted = None
    async for message in iterator:
        if not isinstance(message, StateMessage) or "in_progress" in message.state:
            continue
        file_states = message.state["files"]
        if file_states["a.txt"]["parser_fingerprint"] != old_state["parser_fingerprint"]:
            interrupted = message.state
            break
    await iterator.aclose()

    assert interrupted is not None
    assert interrupted["files"]["b.txt"]["parser_fingerprint"] == old_state["parser_fingerprint"]
    resumed = await collect(migrated, stream, interrupted)
    assert any(item.envelope.metadata["filename"] == "b.txt" for item in records(resumed))


@pytest.mark.asyncio
async def test_changed_file_after_partial_checkpoint_deletes_committed_tail(tmp_path: Path):
    path = tmp_path / "notes.txt"
    path.write_text("replacement", encoding="utf-8")
    source = LocalDocumentsSource(path, stream_names=["notes"])
    stream = (await source.discover())[0]
    parser_fingerprint = source._parser_fingerprint()
    state = {
        "parser_fingerprint": parser_fingerprint,
        "files": {},
        "in_progress": {
            "relative_path": "notes.txt",
            "sha256": "a" * 64,
            "next_index": 5,
            "parser_fingerprint": parser_fingerprint,
        },
    }

    changed = records(await collect(source, stream, state))

    assert [item.envelope.operation for item in changed] == [
        Operation.UPSERT,
        Operation.DELETE,
        Operation.DELETE,
        Operation.DELETE,
        Operation.DELETE,
    ]


@pytest.mark.asyncio
async def test_shrink_and_deleted_files_remove_stale_current_view_records(tmp_path: Path):
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    path = source_dir / "notes.txt"
    path.write_text("many words " * 300, encoding="utf-8")
    state = MemoryStateStore()
    database = tmp_path / "documents.db"

    async def sync() -> None:
        await Pipeline(
            "documents",
            LocalDocumentsSource(
                source_dir,
                stream_names=["documents"],
                text_chunk_chars=256,
            ),
            SQLiteCollection(database),
            state_store=state,
        ).run()

    async def current_count() -> int:
        store = SQLiteCollection(database)
        try:
            return len(await store.query(QueryRequest(limit=1000)))
        finally:
            await store.close()

    await sync()
    assert await current_count() > 1
    path.write_text("one short record", encoding="utf-8")
    await sync()
    assert await current_count() == 1
    path.unlink()
    await sync()
    assert await current_count() == 0


@pytest.mark.asyncio
async def test_deleted_in_progress_file_emits_tombstones(tmp_path: Path):
    source = LocalDocumentsSource(tmp_path, stream_names=["documents"])
    stream = (await source.discover())[0]
    state = {
        "parser_fingerprint": "prior-parser",
        "files": {},
        "in_progress": {
            "relative_path": "removed.txt",
            "sha256": "a" * 64,
            "next_index": 2,
        },
    }

    messages = await collect(source, stream, state)

    deleted = records(messages)
    assert len(deleted) == 2
    assert all(item.envelope.operation is Operation.DELETE for item in deleted)


@pytest.mark.asyncio
async def test_csv_duplicate_headers_and_excel_temporal_values_are_json_safe(tmp_path: Path):
    (tmp_path / "people.csv").write_text("name,name\nAda,Grace\n", encoding="utf-8")
    if documents.load_workbook is None:
        pytest.skip("openpyxl is not installed")
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Events"
    sheet.append(["day", "time", "duration", "created_at"])
    sheet.append(
        [date(2026, 7, 12), time(8, 30), timedelta(minutes=5), datetime(2026, 7, 12, 8, 30)]
    )
    workbook.save(tmp_path / "events.xlsx")

    source = LocalDocumentsSource(tmp_path, stream_names=["tables"], table_batch_rows=1)
    stream = (await source.discover())[0]
    table_records = records(await collect(source, stream))
    payloads = {item.envelope.metadata["filename"]: item.envelope.payload for item in table_records}
    csv_payload = payloads["people.csv"]
    excel_payload = payloads["events.xlsx"]
    assert isinstance(csv_payload, TableBatch)
    assert csv_payload.rows[0] == {"name": "Ada", "name_2": "Grace"}
    assert isinstance(excel_payload, TableBatch)
    assert excel_payload.rows[0] == {
        "day": "2026-07-12T00:00:00",
        "time": "08:30:00",
        "duration": 300.0,
        "created_at": "2026-07-12T08:30:00",
    }

    destination = tmp_path / "temporal.jsonl"
    await Pipeline(
        "temporal",
        LocalDocumentsSource(tmp_path / "events.xlsx", stream_names=["events"]),
        JsonlDestination(destination),
        state_store=MemoryStateStore(),
    ).run()
    assert destination.read_text(encoding="utf-8")


@pytest.mark.asyncio
async def test_word_email_html_and_pdf_emit_bounded_elements(tmp_path: Path, monkeypatch):
    if documents.WordDocument is None:
        pytest.skip("python-docx is not installed")
    from docx import Document

    word = Document()
    word.add_heading("Quarterly report")
    word.add_paragraph("Revenue increased.")
    table = word.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "metric"
    table.rows[0].cells[1].text = "value"
    for index in range(1, 4):
        table.rows[index].cells[0].text = f"metric-{index}"
        table.rows[index].cells[1].text = str(index)
    word.save(tmp_path / "report.docx")

    message = EmailMessage()
    message["Subject"] = "Status"
    message.set_content("Everything is green.")
    (tmp_path / "status.eml").write_bytes(message.as_bytes())
    (tmp_path / "page.html").write_text(
        "<style>hidden</style><h1>Hello</h1><p>from HTML</p>", encoding="utf-8"
    )

    class FakePage:
        def extract_text(self):
            return "PDF page text"

    class FakeReader:
        def __init__(self, _path):
            self.pages = [FakePage()]

    monkeypatch.setattr(documents, "PdfReader", FakeReader)
    (tmp_path / "paper.pdf").write_bytes(b"fake-pdf")

    source = LocalDocumentsSource(
        tmp_path,
        stream_names=["mixed"],
        table_batch_rows=1,
    )
    stream = (await source.discover())[0]
    grouped: dict[str, list] = {}
    for message in records(await collect(source, stream)):
        grouped.setdefault(str(message.envelope.metadata["filename"]), []).append(
            message.envelope.payload
        )

    word_tables = [item for item in grouped["report.docx"] if isinstance(item, TableBatch)]
    assert len(word_tables) == 3
    assert all(len(item.rows) == 1 for item in word_tables)
    assert grouped["status.eml"][0].text == "Everything is green."
    assert grouped["page.html"][0].text == "Hello\nfrom HTML"
    assert grouped["paper.pdf"][0].page_number == 1
    assert any(isinstance(item, DocumentElement) for item in grouped["report.docx"])


@pytest.mark.asyncio
async def test_symlink_root_is_rejected_when_following_is_disabled(tmp_path: Path, monkeypatch):
    root = tmp_path / "documents"
    root.mkdir()
    source = LocalDocumentsSource(root)
    original = Path.is_symlink
    monkeypatch.setattr(Path, "is_symlink", lambda value: value == root or original(value))

    result = await source.check()

    assert result.ok is False
    assert "must not be symlinks" in result.message


@pytest.mark.asyncio
async def test_symlink_swap_before_open_is_rejected(tmp_path: Path, monkeypatch):
    root = tmp_path / "documents"
    root.mkdir()
    path = root / "safe.txt"
    path.write_text("safe", encoding="utf-8")
    secret = tmp_path / "secret.txt"
    secret.write_text("secret", encoding="utf-8")
    real_snapshot = documents._snapshot_file

    def swap_before_open(candidate: Path, **kwargs):
        candidate.unlink()
        try:
            candidate.symlink_to(secret)
        except OSError as exc:
            pytest.skip(f"symlink creation is unavailable: {exc}")
        return real_snapshot(candidate, **kwargs)

    monkeypatch.setattr(documents, "_snapshot_file", swap_before_open)
    source = LocalDocumentsSource(root, stream_names=["documents"])
    stream = (await source.discover())[0]

    with pytest.raises(ConfigurationError, match="symlink|escaped|changed"):
        await collect(source, stream)


@pytest.mark.asyncio
async def test_file_size_limit_is_enforced_during_check(tmp_path: Path):
    (tmp_path / "large.txt").write_text("too large", encoding="utf-8")
    result = await LocalDocumentsSource(tmp_path, max_file_size_bytes=3).check()
    assert result.ok is False
    assert "size limit" in result.message


def test_snapshot_copy_enforces_cumulative_size_limit(tmp_path: Path, monkeypatch):
    path = tmp_path / "growing.txt"
    path.write_bytes(b"0123456789")
    real_fstat = documents.os.fstat

    def small_initial_stat(descriptor: int):
        result = real_fstat(descriptor)
        values = list(result)
        values[6] = 3
        return documents.os.stat_result(values)

    monkeypatch.setattr(documents.os, "fstat", small_initial_stat)
    with (
        pytest.raises(ConfigurationError, match="size limit"),
        documents._snapshot_file(
            path,
            root=path,
            trusted_root=documents._trusted_root(path),
            follow_symlinks=False,
            max_bytes=3,
        ),
    ):
        pass


@pytest.mark.asyncio
async def test_followed_symlink_cycle_terminates_without_duplicate_files(tmp_path: Path):
    root = tmp_path / "documents"
    root.mkdir()
    (root / "notes.txt").write_text("hello", encoding="utf-8")
    try:
        (root / "loop").symlink_to(root, target_is_directory=True)
    except OSError as exc:
        pytest.skip(f"directory symlink creation is unavailable: {exc}")
    source = LocalDocumentsSource(root, stream_names=["documents"], follow_symlinks=True)
    stream = (await source.discover())[0]

    assert len(records(await collect(source, stream))) == 1


@pytest.mark.asyncio
async def test_direct_file_deletion_reconciles_current_view(tmp_path: Path):
    path = tmp_path / "notes.txt"
    path.write_text("hello", encoding="utf-8")
    state = MemoryStateStore()
    database = tmp_path / "direct.db"

    async def sync() -> None:
        await Pipeline(
            "direct",
            LocalDocumentsSource(path, stream_names=["notes"]),
            SQLiteCollection(database),
            state_store=state,
        ).run()

    await sync()
    path.unlink()
    await sync()
    store = SQLiteCollection(database)
    try:
        assert len(await store.query(QueryRequest(limit=10))) == 0
    finally:
        await store.close()


@pytest.mark.asyncio
async def test_missing_optional_parser_is_reported_by_check(tmp_path: Path, monkeypatch):
    (tmp_path / "paper.pdf").write_bytes(b"pdf")
    monkeypatch.setattr(documents, "PdfReader", None)
    result = await LocalDocumentsSource(tmp_path).check()
    assert result.ok is False
    assert "ingestion-graph[documents]" in result.message


@pytest.mark.asyncio
async def test_document_source_runs_in_embedded_pipeline(tmp_path: Path):
    source_path = tmp_path / "source"
    source_path.mkdir()
    (source_path / "notes.md").write_text("Reusable SDK ingestion", encoding="utf-8")
    destination_path = tmp_path / "records.jsonl"
    result = await Pipeline(
        "documents",
        LocalDocumentsSource(source_path, stream_names=["documents"]),
        JsonlDestination(destination_path),
        state_store=MemoryStateStore(),
    ).run()
    assert result.records_written == 1
    assert "Reusable SDK ingestion" in destination_path.read_text(encoding="utf-8")
